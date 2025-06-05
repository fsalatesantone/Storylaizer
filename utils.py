import streamlit as st
import time
import pandas as pd
import numpy as np
import io
from docx import Document
import markdown2
from html2docx import html2docx
import re
from typing import Any
import traceback


def reset_conversation():
    # Nuovo ID sessione per forzare ricaricamento
    st.session_state.session_id = str(time.time())
    
    # Reset variabili principali
    keys_to_reset = ["chat_history1", "chat_history2", "chat_history3"
                     , "file_loaded1", "uploaded_file1"
                     , "file_loaded2", "uploaded_file2"
                     , "dataframe", "dataframe_report", "data_metadata", "data_errors"]
    for key in keys_to_reset:
        if key in st.session_state:
            del st.session_state[key]
            
def init_session_state():
    if "chat_history1" not in st.session_state:
        st.session_state.chat_history1 = []
    if "chat_history2" not in st.session_state:
        st.session_state.chat_history2 = []
    if "chat_history3" not in st.session_state:
        st.session_state.chat_history3 = []

    if "conversation_started1" not in st.session_state:
        st.session_state.conversation_started1 = False
    if "conversation_started2" not in st.session_state:
        st.session_state.conversation_started2 = False
    if "conversation_started3" not in st.session_state:
        st.session_state.conversation_started3 = False

    if "pending_user_message1" not in st.session_state:
        st.session_state.pending_user_message1 = None
    if "pending_user_message2" not in st.session_state:
        st.session_state.pending_user_message2 = None
    if "pending_user_message3" not in st.session_state:
        st.session_state.pending_user_message3 = None

    if "selected_mode" not in st.session_state:
        st.session_state.selected_mode = None
    if "file_loaded1" not in st.session_state:
        st.session_state.file_loaded1 = False
    if "uploaded_file1" not in st.session_state:
        st.session_state.uploaded_file1 = None
    if "file_loaded2" not in st.session_state:
        st.session_state.file_loaded2 = False
    if "uploaded_file2" not in st.session_state:
        st.session_state.uploaded_file2 = None
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(time.time())

def export_chat(format_type, chat_history):
    if chat_history == []:
        st.warning("Non ci sono messaggi da esportare.")
        return None

    if format_type == "xlsx":
        conv_rows = []   # Lista di dict: ogni dict rappresenta una riga in Conversazione

        for msg in chat_history:
            role = "Utente" if msg["role"] == "user" else "Assistente"
            content = msg["content"] or ""
            lines = content.splitlines()

            # Se il messaggio è vuoto, lo tratto come riga unica di testo vuoto
            if not lines:
                conv_rows.append({"Ruolo": role, "Messaggio": ""})
                continue

            i = 0
            # Scorro tutte le righe di questo messaggio
            while i < len(lines):
                # Caso 1: incontro un possibile blocco-tabella Markdown
                # Deve esserci almeno una riga per header e una riga per separatore
                if (
                    i + 1 < len(lines)
                    and re.match(r"^\|(.+\|)+\s*$", lines[i])
                    and re.match(r"^\|\s*[-:]+\s*(\|\s*[-:]+\s*)+\|\s*$", lines[i+1])
                ):
                    # Trovo la tabella a partire da lines[i]
                    headers = [h.strip() for h in lines[i].strip("|").split("|")]

                    # Raccolgo tutte le righe dati fino a che iniziano con ‘|’
                    data_rows = []
                    j = i + 2
                    while j < len(lines) and lines[j].strip().startswith("|"):
                        data_rows.append([c.strip() for c in lines[j].strip("|").split("|")])
                        j += 1

                    # 1) Se prima della tabella c’erano righe di testo “normale”, le inserisco in una riga separata
                    #    (solo se i > partenza di questo segmento)
                    #    Qui però, siccome stiamo entrando solo quando troviamo la tabella, 
                    #    eventuali righe di testo provenienti da cicli precedenti sono già state gestite.

                    # 2) Inserisco la riga “header” della tabella nel foglio Excel:
                    row_header = {"Ruolo": role}
                    for idx, hdr in enumerate(headers, start=1):
                        row_header[f"Col{idx}"] = hdr
                    conv_rows.append(row_header)

                    # 3) Inserisco le righe dati, con colonna “Ruolo” vuota
                    for data in data_rows:
                        row_data = {"Ruolo": ""}
                        for idx, val in enumerate(data, start=1):
                            row_data[f"Col{idx}"] = val
                        conv_rows.append(row_data)

                    # 4) Salto tutto il blocco tabella
                    i = j

                else:
                    # Caso 2: questa riga NON fa parte di una tabella markdown.
                    # Resto in “modalità testo normale” finché non trovo l’inizio di una tabella
                    testo_accumulato = []
                    while i < len(lines):
                        # Se la riga corrente è inizio di tabella, interrompo il “testo normale”
                        if (
                            i + 1 < len(lines)
                            and re.match(r"^\|(.+\|)+\s*$", lines[i])
                            and re.match(r"^\|\s*[-:]+\s*(\|\s*[-:]+\s*)+\|\s*$", lines[i+1])
                        ):
                            break
                        # Altrimenti accumulo questa riga in testo_accumulato
                        testo_accumulato.append(lines[i])
                        i += 1

                    # Inserisco tutto il blocco “testo normale” in un’unica riga nel DataFrame
                    testo_unico = "\n".join(testo_accumulato).strip()
                    conv_rows.append({
                        "Ruolo": role,
                        "Messaggio": testo_unico
                    })
                    # Torno nel loop principale (che controllerà di nuovo se ora ci sono tabelle)
            # Fine while i < len(lines)

        # A questo punto conv_rows contiene sia:
        # - righe con chiavi “Ruolo” + “Messaggio” (messaggi normali)
        # - righe con chiavi “Ruolo” vuoto o “Ruolo” + Col1, Col2, ... (per le tabelle)

        # Costruisco il DataFrame rettangolare, unificando tutte le colonne possibili
        df_conv = pd.DataFrame(conv_rows).fillna("")

        # Compongo l’ordine definitivo delle colonne:
        #   1) “Ruolo”
        #   2) “Messaggio” (se presente)
        #   3) tutte le “Col1”, “Col2”, ... in ordine numerico
        tutte_colonne = ["Ruolo"]
        if "Messaggio" in df_conv.columns:
            tutte_colonne.append("Messaggio")

        col_tabella = sorted(
            [c for c in df_conv.columns if c.startswith("Col")],
            key=lambda x: int(x.replace("Col", ""))
        )
        tutte_colonne.extend(col_tabella)

        # Ricreo il DataFrame con solo queste colonne, riempiendo eventuali buchi con stringhe vuote
        df_final = df_conv.reindex(columns=tutte_colonne).fillna("")

        # Scrittura su Excel in un unico foglio “Conversazione”
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df_final.to_excel(writer, index=False, sheet_name="Conversazione")

        return (
            output.getvalue(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "conversazione.xlsx"
        )

    elif format_type == "txt":
        content = ""
        for msg in chat_history:
            prefix = "Utente: " if msg["role"] == "user" else "Assistente: "
            content += f"{prefix}{msg['content']}\n\n"
        return content.encode(), "text/plain", "conversazione.txt"

    elif format_type == "docx":
        try:
            final_doc = Document()

            for idx, msg in enumerate(chat_history, start=1):
                text = msg.get("content", "").strip()
                if not text:
                    continue

                role = "Utente" if msg["role"] == "user" else "Assistente"
                markdown_text = f"**{role}:**\n\n{text}"

                # Provo a convertire tutto in HTML (inclusi i blocchi tabella)
                html = markdown2.markdown(markdown_text, extras=["tables", "fenced-code-blocks"])
                try:
                    temp_stream = html2docx(html, title=f"msg_{idx}")
                    from docx import Document as _Doc
                    temp_doc = _Doc(io.BytesIO(temp_stream.getvalue()))
                    for element in temp_doc.element.body:
                        final_doc.element.body.append(element)
                except Exception:
                    # Fallback manual per tabelle Markdown
                    lines = markdown_text.splitlines()
                    if (len(lines) >= 3
                        and re.match(r"^\|.+\|$", lines[0])
                        and re.match(r"^\|\s*[-:]+\s*(\|\s*[-:]+\s*)+\|$", lines[1])):

                        headers = [h.strip() for h in lines[0].strip("|").split("|")]
                        data_rows = []
                        for row in lines[2:]:
                            if not row.strip().startswith("|"):
                                break
                            data_rows.append([c.strip() for c in row.strip("|").split("|")])

                        table = final_doc.add_table(rows=1 + len(data_rows), cols=len(headers))


                        for i, hdr in enumerate(headers):
                            table.rows[0].cells[i].text = hdr
                        for r, row in enumerate(data_rows, start=1):
                            for c, val in enumerate(row):
                                table.rows[r].cells[c].text = val
                        final_doc.add_paragraph("")  # spazio dopo la tabella
                    else:
                        p = final_doc.add_paragraph()
                        p.add_run(markdown_text)

                final_doc.add_paragraph("")  # spazio tra i messaggi

            output = io.BytesIO()
            final_doc.save(output)
            return (
                output.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "conversazione.docx"
            )
        except ImportError:
            st.error(
                "Per esportare in DOCX con Markdown servono: "
                "python-docx, markdown2, html2docx"
            )
            return None


# formatta l'output come tabella Markdown
def _format_list_of_dicts_as_markdown(data):
    if not data or not isinstance(data, list):
        return data
    if not all(isinstance(row, dict) for row in data):
        return data

    keys = list({key for row in data for key in row})
    if not keys:
        return data

    header = "| " + " | ".join(keys) + " |"
    separator = "| " + " | ".join(["---"] * len(keys)) + " |"
    rows = []
    for row in data:
        line = "| " + " | ".join(str(row.get(k, "")) for k in keys) + " |"
        rows.append(line)

    return "\n".join([header, separator] + rows)


# Funzione sandboxed per eseguire codice su df
def execute_code(code: str, df: pd.DataFrame) -> Any:
    """
    Esegue in sandbox un blocco di codice Python che utilizza la variabile `df`.
    Si attende che venga definita in esso una variabile `result`.
    Restituisce sempre JSON serializzabile senza includere indici.
    L'accesso alle colonne è case-insensitive e supporta anche accesso fuzzy.
    """
    
    class CIDataFrame(pd.DataFrame):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            # Crea mappatura case-insensitive e fuzzy delle colonne
            #self._col_map = {}
            object.__setattr__(self, "_col_map", {})
            for col in self.columns:
                col_lower = str(col).lower()
                self._col_map[col_lower] = col
                # Aggiungi anche versioni plurali/singolari comuni
                if col_lower.endswith('e'):
                    self._col_map[col_lower[:-1] + 'i'] = col  # regione -> regioni
                elif col_lower.endswith('a'):
                    self._col_map[col_lower[:-1] + 'e'] = col  # persona -> persone
                elif col_lower.endswith('o'):
                    self._col_map[col_lower[:-1] + 'i'] = col  # prodotto -> prodotti
                elif col_lower.endswith('i'):
                    self._col_map[col_lower[:-1] + 'e'] = col  # regioni -> regione
                    self._col_map[col_lower[:-1] + 'o'] = col  # prodotti -> prodotto
                elif col_lower.endswith('e') and len(col_lower) > 3:
                    self._col_map[col_lower[:-1] + 'a'] = col  # persone -> persona

        @property
        def _constructor(self):
            return CIDataFrame
        
        def _find_column(self, key):
            """Trova la colonna corrispondente ignorando case e con fuzzy matching"""
            if not isinstance(key, str):
                return key
            
            key_lower = key.lower()
            original_key = key  # Salva per debugging
            
            # 1. Cerca match esatto case-insensitive
            if key_lower in self._col_map:
                found = self._col_map[key_lower]
                print(f"Column match: '{original_key}' -> '{found}' (exact case-insensitive)")
                return found
            
            # 2. Normalizza la chiave (rimuovi spazi, underscore, etc.)
            key_normalized = key_lower.replace(' ', '').replace('_', '').replace('-', '')
            
            # 3. Cerca match con normalizzazione
            for col in self.columns:
                col_normalized = str(col).lower().replace(' ', '').replace('_', '').replace('-', '')
                if key_normalized == col_normalized:
                    print(f"Column match: '{original_key}' -> '{col}' (normalized)")
                    return col
            
            # 4. Cerca match parziale (contiene) - con e senza normalizzazione
            for col in self.columns:
                col_lower = str(col).lower()
                col_normalized = col_lower.replace(' ', '').replace('_', '').replace('-', '')
                
                # Match parziale normale
                if (key_lower in col_lower or col_lower in key_lower or
                    # Match parziale normalizzato
                    key_normalized in col_normalized or col_normalized in key_normalized):
                    print(f"Column match: '{original_key}' -> '{col}' (partial match)")
                    return col
            
            # 5. Match fuzzy per parole chiave comuni
            key_words = set(key_lower.split())
            for col in self.columns:
                col_words = set(str(col).lower().split())
                # Se almeno il 50% delle parole corrispondono
                if len(key_words & col_words) >= len(key_words) * 0.5:
                    print(f"Column match: '{original_key}' -> '{col}' (fuzzy word match)")
                    return col
            
            # 6. Se non trova niente, restituisce la chiave originale
            print(f"Column match: '{original_key}' -> '{key}' (no match found, using original)")
            print(f"Available columns: {list(self.columns)}")
            return key

        def __getitem__(self, key):
            if isinstance(key, str):
                matched_key = self._find_column(key)
                return super().__getitem__(matched_key)
            elif isinstance(key, list):
                # Gestisci liste di colonne
                matched_keys = [self._find_column(k) if isinstance(k, str) else k for k in key]
                return super().__getitem__(matched_keys)
            return super().__getitem__(key)
        
        def __getattr__(self, name):
            # Gestisce l'accesso df.colonna
            if name.startswith('_') or name in ['columns', 'index', 'values']:
                return super().__getattribute__(name)
            
            try:
                matched_name = self._find_column(name)
                if matched_name in self.columns:
                    return self[matched_name]
                return super().__getattribute__(name)
            except AttributeError:
                # Se non è una colonna, prova l'attributo normale
                return super().__getattribute__(name)
        
        def groupby(self, by, **kwargs):
            # Gestisce groupby con nomi case-insensitive
            if isinstance(by, str):
                by = self._find_column(by)
            elif isinstance(by, list):
                by = [self._find_column(col) if isinstance(col, str) else col for col in by]
            return super().groupby(by, **kwargs)
        
        def sort_values(self, by, **kwargs):
            # Gestisce sort_values con nomi case-insensitive
            if isinstance(by, str):
                by = self._find_column(by)
            elif isinstance(by, list):
                by = [self._find_column(col) if isinstance(col, str) else col for col in by]
            return super().sort_values(by, **kwargs)
        
        def drop(self, labels, **kwargs):
            # Gestisce drop con nomi case-insensitive
            if isinstance(labels, str):
                labels = self._find_column(labels)
            elif isinstance(labels, list):
                labels = [self._find_column(col) if isinstance(col, str) else col for col in labels]
            return super().drop(labels, **kwargs)

    # Crea proxy CI
    df_ci = CIDataFrame(df)
    
    # Prepara builtins limitati
    import builtins as _builtins_module
    allowed_builtins = {
        name: getattr(_builtins_module, name) 
        for name in ['len','sum','min','max','sorted','list','dict','float','int','str','round', 
                     'abs', 'filter', 'map', 'zip', 'enumerate', 'range', 'print', 'set', 'tuple']
    }
    
    safe_globals = {"__builtins__": allowed_builtins, "df": df_ci, "pd": pd, "np": np}
    local_vars = {}
    
    try:
        # Pulisci il codice da formattazione markdown
        # Rimuovi ```python e ``` all'inizio e fine
        code = re.sub(r'^```(?:python)?\s*\n', '', code, flags=re.MULTILINE)
        code = re.sub(r'\n```\s*$', '', code, flags=re.MULTILINE)
        
        # Rimuovi asterischi attorno alle parole chiave Python
        keywords = ['import', 'from', 'def', 'class', 'if', 'else', 'elif', 'for', 'while', 'try', 'except', 'finally', 'with', 'return', 'yield', 'lambda', 'and', 'or', 'not', 'in', 'is']
        for keyword in keywords:
            # Pattern per **keyword** -> keyword
            code = re.sub(rf'\*\*{keyword}\*\*', keyword, code, flags=re.IGNORECASE)
            # Pattern per *keyword* -> keyword  
            code = re.sub(rf'(?<!\*)\*{keyword}\*(?!\*)', keyword, code, flags=re.IGNORECASE)
        
        # Rimuovi altri pattern markdown comuni
        code = re.sub(r'\*\*(.*?)\*\*', r'\1', code)  # **testo** -> testo
        code = re.sub(r'(?<!\*)\*(.*?)\*(?!\*)', r'\1', code)  # *testo* -> testo (ma non **)
        
        # Rimuovi commenti markdown che iniziano con molti #
        lines = code.split('\n')
        cleaned_lines = []
        for line in lines:
            # Se la riga inizia con più di 3 #, probabilmente è un titolo markdown
            if re.match(r'^#{4,}', line.strip()):
                continue
            cleaned_lines.append(line)

        # Intercetta import permessi (solo numpy e pandas)
        import_lines = []
        remaining_lines = []
        for line in code.splitlines():
            stripped = line.strip()
            if re.match(r"^import\s+(numpy\s+as\s+np|pandas\s+as\s+pd)$", stripped):
                # Import permesso → non lo includiamo nel codice, ma lo "simuliamo"
                import_lines.append(stripped)
            elif re.match(r"^\s*(import|from\s+\w+\s+import)", stripped):
                return {
                    "error": f"Import vietato: '{stripped}'. Solo 'import numpy as np' o 'import pandas as pd' sono ammessi.",
                    "code": code
                }
            else:
                remaining_lines.append(line)

        # Sovrascrivi il codice senza le righe di import
        code = '\n'.join(remaining_lines)

        # Prova prima ad eseguire il codice normale
        exec(code, safe_globals, local_vars)
        raw = local_vars.get("result")
        
        if raw is None:
            # Cerca automaticamente variabili che potrebbero essere il risultato
            result_candidates = []
            for var_name, var_value in local_vars.items():
                if not var_name.startswith('_'):
                    # Priorità a variabili DataFrame/Series o con nomi suggestivi
                    if (isinstance(var_value, (pd.DataFrame, pd.Series)) or 
                        any(keyword in var_name.lower() for keyword in ['result', 'top', 'bottom', 'final', 'output', 'answer'])):
                        result_candidates.append((var_name, var_value))
            
            # Se c'è un solo candidato ovvio, usalo
            if len(result_candidates) == 1:
                raw = result_candidates[0][1]
            # Se ci sono più candidati, prendi l'ultimo definito che è DataFrame/Series
            elif len(result_candidates) > 1:
                for var_name, var_value in reversed(list(local_vars.items())):
                    if isinstance(var_value, (pd.DataFrame, pd.Series)) and not var_name.startswith('_'):
                        raw = var_value
                        break
        
        if raw is None:
            lines = code.strip().split('\n')
            if lines:
                last_line = lines[-1].strip()
                preceding_lines = lines[:-1]

                # Se l'ultima riga sembra una "espressione", proviamo ad estrarre il valore
                if (
                    last_line and
                    not last_line.startswith('#') and
                    not any(last_line.startswith(stmt) for stmt in ['if ', 'for ', 'while ', 'def ', 'class ', 'import ', 'from '])
                ):
                    try:
                        # Esegui tutto il codice precedente
                        exec('\n'.join(preceding_lines), safe_globals, local_vars)

                        # Valuta l'ultima riga come espressione
                        raw = eval(last_line, safe_globals, local_vars)
                    except Exception:
                        # Se fallisce, prova ad assegnarlo come `result = ...`
                        try:
                            modified_code = '\n'.join(preceding_lines + [f'result = {last_line}'])
                            local_vars_retry = {}
                            exec(modified_code, safe_globals, local_vars_retry)
                            raw = local_vars_retry.get("result")
                        except:
                            pass
        
        if raw is None:
            # Debug: mostra tutte le variabili definite
            defined_vars = [k for k in local_vars.keys() if not k.startswith('_')]
            return {
                "error": "La variabile 'result' non è stata definita.",
                "debug_info": {
                    "defined_variables": defined_vars,
                    "executed_code": code,
                    "suggestion": "Assicurati che l'ultima riga sia 'result = ...' oppure che definisci esplicitamente 'result'",
                    "dataframe_info": {
                        "shape": df_ci.shape,
                        "columns": list(df_ci.columns),
                        "sample_data": df_ci.head(3).to_dict('records') if len(df_ci) > 0 else []
                    }
                }
            }
        
        # Rimuovi indici e serializza DataFrame/Series
        if isinstance(raw, pd.DataFrame):
            result_data = raw.reset_index(drop=True).to_dict(orient="records")
            return _format_list_of_dicts_as_markdown(result_data)

        if isinstance(raw, pd.Series):
            result_data = raw.reset_index(drop=True).to_frame(name="valore").to_dict(orient="records")
            return _format_list_of_dicts_as_markdown(result_data)

        # Se è una lista di dizionari, prova comunque a formattare come Markdown
        if isinstance(raw, list) and all(isinstance(r, dict) for r in raw):
            return _format_list_of_dicts_as_markdown(raw)

        # Se è una lista/dict vuota, aggiungi info di debug
        if isinstance(raw, (list, dict)) and not raw:
            print(f"Warning: Result is empty {type(raw).__name__}")

        return raw
        
    except Exception as e:
        return {"error": str(e), "traceback": traceback.format_exc(), "code": code}